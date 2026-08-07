from pathlib import Path
from pypdf import PdfReader
from docx import Document as DocxDocument
import os

class DocumentParser:
    """Parse PDF and DOCX documents and extract text"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
    
    def parse_pdf(self, file_path):
        """Extract text from PDF"""
        try:
            reader = PdfReader(file_path)
            text = ""
            metadata = {
                'filename': os.path.basename(file_path),
                'pages': len(reader.pages),
                'format': 'PDF'
            }
            
            for page_num, page in enumerate(reader.pages):
                text += f"\n--- Page {page_num + 1} ---\n"
                text += page.extract_text()
            
            return text, metadata
        except Exception as e:
            print(f"Error parsing PDF {file_path}: {str(e)}")
            return None, None
    
    def parse_docx(self, file_path):
        """Extract text from DOCX"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            metadata = {
                'filename': os.path.basename(file_path),
                'paragraphs': len(doc.paragraphs),
                'format': 'DOCX'
            }
            
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            return text, metadata
        except Exception as e:
            print(f"Error parsing DOCX {file_path}: {str(e)}")
            return None, None
    
    def parse_txt(self, file_path):
        """Extract text from TXT"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            metadata = {
                'filename': os.path.basename(file_path),
                'format': 'TXT'
            }
            return text, metadata
        except Exception as e:
            print(f"Error parsing TXT {file_path}: {str(e)}")
            return None, None
    
    def parse_document(self, file_path):
        """Auto-detect format and parse"""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return self.parse_pdf(file_path)
        elif file_ext == '.docx':
            return self.parse_docx(file_path)
        elif file_ext == '.txt':
            return self.parse_txt(file_path)
        else:
            print(f"Unsupported format: {file_ext}")
            return None, None

# Initialize parser
parser = DocumentParser()
print("✓ Document Parser initialized")